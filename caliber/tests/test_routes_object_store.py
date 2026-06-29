"""Object Store console routes — boto3 against an S3 backend, mocked with moto.

``moto.mock_aws`` patches botocore's API-call layer, so the route's boto3 client
(built with the MinIO endpoint_url) is intercepted regardless of the endpoint.
"""

from __future__ import annotations

import io

import pytest
from moto import mock_aws
from starlette.testclient import TestClient

PREFIX = "/ajax-api/2.0/mlflow/caliber"
OS = PREFIX + "/object-store"


@pytest.fixture(autouse=True)
def _minio_creds(monkeypatch: pytest.MonkeyPatch) -> None:
    # resolve_secret(object_store_access_key_source) reads these env vars.
    monkeypatch.setenv("MINIO_ROOT_USER", "testing")
    monkeypatch.setenv("MINIO_ROOT_PASSWORD", "testing")


def _wire_moto(client: TestClient) -> None:
    """Pre-build the route's boto3 client under the active moto mock.

    moto patches the AWS-default S3 endpoint, not arbitrary endpoints like the
    configured MinIO :9000 URL — so we inject a default-endpoint client (built
    inside ``mock_aws``) for the route to use, exercising the handlers against
    moto's in-memory S3.
    """
    import boto3

    client.app.state.object_store_client = boto3.client("s3", region_name="us-east-1")


@mock_aws
def test_status_buckets_objects_full_crud(client: TestClient) -> None:
    _wire_moto(client)
    # status — reachable, no buckets yet
    status = client.get(f"{OS}/status")
    assert status.status_code == 200
    assert status.json()["data"]["connected"] is True

    # create + list buckets
    cb = client.post(f"{OS}/buckets", json={"name": "reports"})
    assert cb.status_code == 201, (cb.status_code, cb.text)
    assert "reports" in [b["name"] for b in client.get(f"{OS}/buckets").json()["data"]]

    # upload an object under a prefix
    files = {"file": ("kg.json", io.BytesIO(b'{"k":1}'), "application/json")}
    up = client.post(f"{OS}/buckets/reports/objects", data={"prefix": "2026/"}, files=files)
    assert up.status_code == 201, up.text
    assert up.json()["data"]["key"] == "2026/kg.json"

    # list at root → folder "2026/"; list inside → the object
    root = client.get(f"{OS}/buckets/reports/objects").json()["data"]
    assert "2026/" in root["prefixes"]
    inside = client.get(f"{OS}/buckets/reports/objects", params={"prefix": "2026/"}).json()["data"]
    obj = next(o for o in inside["objects"] if o["key"] == "2026/kg.json")
    assert obj["size"] == 7
    assert obj["created_at"] is not None
    assert obj["last_modified"] is not None

    # download streams the bytes back
    dl = client.get(f"{OS}/buckets/reports/object", params={"key": "2026/kg.json"})
    assert dl.status_code == 200 and dl.content == b'{"k":1}'

    # delete object, then bucket
    assert (
        client.delete(f"{OS}/buckets/reports/object", params={"key": "2026/kg.json"}).status_code
        == 204
    )
    assert client.delete(f"{OS}/buckets/reports").status_code == 204
    assert "reports" not in [b["name"] for b in client.get(f"{OS}/buckets").json()["data"]]


@mock_aws
def test_create_bucket_rejects_invalid_name(client: TestClient) -> None:
    assert client.post(f"{OS}/buckets", json={"name": "A_B C"}).status_code == 400


@mock_aws
def test_download_missing_object_404(client: TestClient) -> None:
    _wire_moto(client)
    client.post(f"{OS}/buckets", json={"name": "bkt1"})
    r = client.get(f"{OS}/buckets/bkt1/object", params={"key": "nope.txt"})
    assert r.status_code == 404


@mock_aws
def test_mutations_require_admin(client: TestClient) -> None:
    _wire_moto(client)
    viewer = {"X-CALIBER-User": "@viewer"}
    assert client.post(f"{OS}/buckets", json={"name": "x"}, headers=viewer).status_code == 403
    assert client.delete(f"{OS}/buckets/x", headers=viewer).status_code == 403
    # reads are allowed for any authenticated user
    assert client.get(f"{OS}/buckets", headers=viewer).status_code == 200


def _put(client: TestClient, bucket: str, key: str) -> None:
    client.post(
        f"{OS}/buckets/{bucket}/objects",
        data={"key": key},
        files={"file": (key.rsplit("/", 1)[-1], io.BytesIO(b"x"), "text/plain")},
    )


@mock_aws
def test_create_folder_and_recursive_list(client: TestClient) -> None:
    _wire_moto(client)
    client.post(f"{OS}/buckets", json={"name": "fld"})
    # create a top-level folder → a "{name}/" marker that surfaces as a prefix
    mk = client.post(f"{OS}/buckets/fld/folders", json={"prefix": "", "name": "docs"})
    assert mk.status_code == 201, mk.text
    assert mk.json()["data"]["prefix"] == "docs/"
    root = client.get(f"{OS}/buckets/fld/objects").json()["data"]
    assert "docs/" in root["prefixes"]
    assert root["objects"] == []  # the 0-byte marker is filtered out, not shown as a file

    # nested folder + files, then compare one-level vs recursive listing
    client.post(f"{OS}/buckets/fld/folders", json={"prefix": "docs/", "name": "2026"})
    _put(client, "fld", "docs/2026/a.txt")
    _put(client, "fld", "docs/2026/b.txt")
    lvl = client.get(f"{OS}/buckets/fld/objects", params={"prefix": "docs/"}).json()["data"]
    assert "docs/2026/" in lvl["prefixes"] and lvl["objects"] == []
    rec = client.get(f"{OS}/buckets/fld/objects", params={"recursive": "true"}).json()["data"]
    keys = {o["key"] for o in rec["objects"]}
    assert {"docs/2026/a.txt", "docs/2026/b.txt"} <= keys
    assert not any(k.endswith("/") for k in keys)  # no folder markers in a flat listing


@mock_aws
def test_preview_object_supports_text_and_binary(client: TestClient) -> None:
    _wire_moto(client)
    client.post(f"{OS}/buckets", json={"name": "pvw"})
    client.post(
        f"{OS}/buckets/pvw/objects",
        data={"key": "service/2026/log.jsonl"},
        files={"file": ("log.jsonl", io.BytesIO(b'{"ok":1}\n{"ok":2}\n'), "application/x-ndjson")},
    )
    text_preview = client.get(
        f"{OS}/buckets/pvw/object/preview",
        params={"key": "service/2026/log.jsonl", "max_bytes": "8"},
    )
    assert text_preview.status_code == 200, text_preview.text
    text_data = text_preview.json()["data"]
    assert text_data["is_text"] is True
    assert text_data["truncated"] is True
    assert text_data["text"].startswith('{"ok":1}')
    assert text_data["content_type"] == "application/x-ndjson"
    assert text_data["created_at"] is not None

    client.post(
        f"{OS}/buckets/pvw/objects",
        data={"key": "service/2026/blob.bin"},
        files={"file": ("blob.bin", io.BytesIO(b"\x00\xff\x10\x11"), "application/octet-stream")},
    )
    bin_preview = client.get(
        f"{OS}/buckets/pvw/object/preview",
        params={"key": "service/2026/blob.bin"},
    )
    assert bin_preview.status_code == 200, bin_preview.text
    bin_data = bin_preview.json()["data"]
    assert bin_data["is_text"] is False
    assert bin_data["text"] is None
    assert bin_data["preview_bytes"] == 4
    assert bin_data["created_at"] is not None


@mock_aws
def test_preview_object_rejects_invalid_max_bytes(client: TestClient) -> None:
    _wire_moto(client)
    client.post(f"{OS}/buckets", json={"name": "pvw2"})
    _put(client, "pvw2", "a.txt")
    bad = client.get(f"{OS}/buckets/pvw2/object/preview", params={"key": "a.txt", "max_bytes": "0"})
    assert bad.status_code == 400


def _put_bytes(client: TestClient, bucket: str, key: str, data: bytes, content_type: str) -> None:
    client.post(
        f"{OS}/buckets/{bucket}/objects",
        data={"key": key},
        files={"file": (key.rsplit("/", 1)[-1], io.BytesIO(data), content_type)},
    )


@mock_aws
def test_download_inline_media_type_and_range(client: TestClient) -> None:
    _wire_moto(client)
    client.post(f"{OS}/buckets", json={"name": "media"})
    _put_bytes(client, "media", "clip.mp4", b"0123456789", "application/octet-stream")

    # Inline disposition maps the extension to a streamable media type.
    inline = client.get(
        f"{OS}/buckets/media/object", params={"key": "clip.mp4", "disposition": "inline"}
    )
    assert inline.status_code == 200, inline.text
    assert inline.headers["content-type"].startswith("video/mp4")
    assert inline.headers.get("accept-ranges") == "bytes"

    # A Range request returns the matching slice as a 206 with Content-Range.
    ranged = client.get(
        f"{OS}/buckets/media/object",
        params={"key": "clip.mp4", "disposition": "inline"},
        headers={"Range": "bytes=0-3"},
    )
    assert ranged.status_code == 206, ranged.text
    assert ranged.content == b"0123"
    assert ranged.headers["content-range"].startswith("bytes 0-3/")


@mock_aws
def test_extract_word_document_to_text(client: TestClient) -> None:
    import docx

    _wire_moto(client)
    client.post(f"{OS}/buckets", json={"name": "ext"})
    buf = io.BytesIO()
    document = docx.Document()
    document.add_paragraph("Quarterly memo body.")
    document.add_paragraph("Second paragraph.")
    document.save(buf)
    _put_bytes(client, "ext", "memo.docx", buf.getvalue(), "application/octet-stream")

    res = client.get(f"{OS}/buckets/ext/object/extract", params={"key": "memo.docx"})
    assert res.status_code == 200, res.text
    data = res.json()["data"]
    assert data["kind"] == "document"
    assert "Quarterly memo body." in data["text"]
    assert "Second paragraph." in data["text"]


@mock_aws
def test_extract_excel_workbook_to_sheets(client: TestClient) -> None:
    import openpyxl

    _wire_moto(client)
    client.post(f"{OS}/buckets", json={"name": "ext2"})
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Q1"
    ws.append(["Region", "Total"])
    ws.append(["EMEA", 42])
    buf = io.BytesIO()
    wb.save(buf)
    _put_bytes(client, "ext2", "data.xlsx", buf.getvalue(), "application/octet-stream")

    res = client.get(f"{OS}/buckets/ext2/object/extract", params={"key": "data.xlsx"})
    assert res.status_code == 200, res.text
    data = res.json()["data"]
    assert data["kind"] == "sheet"
    assert data["sheets"][0]["name"] == "Q1"
    rows = data["sheets"][0]["rows"]
    assert rows[0] == ["Region", "Total"]
    assert rows[1] == ["EMEA", "42"]


@mock_aws
def test_extract_unsupported_formats(client: TestClient) -> None:
    _wire_moto(client)
    client.post(f"{OS}/buckets", json={"name": "ext3"})
    _put_bytes(client, "ext3", "legacy.doc", b"\xd0\xcf\x11\xe0junk", "application/msword")
    _put_bytes(client, "ext3", "plain.txt", b"just text", "text/plain")

    legacy = client.get(f"{OS}/buckets/ext3/object/extract", params={"key": "legacy.doc"})
    assert legacy.status_code == 200, legacy.text
    assert legacy.json()["data"]["kind"] == "unsupported"
    assert "Legacy" in legacy.json()["data"]["error"]

    plain = client.get(f"{OS}/buckets/ext3/object/extract", params={"key": "plain.txt"})
    assert plain.json()["data"]["kind"] == "unsupported"


@mock_aws
def test_create_folder_rejects_bad_name_and_duplicate(client: TestClient) -> None:
    _wire_moto(client)
    client.post(f"{OS}/buckets", json={"name": "fld2"})
    assert client.post(f"{OS}/buckets/fld2/folders", json={"name": "a/b"}).status_code == 400
    assert client.post(f"{OS}/buckets/fld2/folders", json={"name": ""}).status_code == 400
    assert client.post(f"{OS}/buckets/fld2/folders", json={"name": "dup"}).status_code == 201
    assert client.post(f"{OS}/buckets/fld2/folders", json={"name": "dup"}).status_code == 409


@mock_aws
def test_batch_delete_keys_and_prefix(client: TestClient) -> None:
    _wire_moto(client)
    client.post(f"{OS}/buckets", json={"name": "bdx"})
    for key in ("top.txt", "f/1.txt", "f/2.txt", "f/sub/3.txt"):
        _put(client, "bdx", key)
    # explicit keys
    r = client.post(f"{OS}/buckets/bdx/objects/delete", json={"keys": ["top.txt"]})
    assert r.status_code == 200 and r.json()["data"]["deleted"] == 1
    # whole "f/" folder (recursive prefix delete)
    r2 = client.post(f"{OS}/buckets/bdx/objects/delete", json={"prefix": "f/"})
    assert r2.status_code == 200 and r2.json()["data"]["deleted"] == 3
    assert (
        client.get(f"{OS}/buckets/bdx/objects", params={"recursive": "true"}).json()["data"][
            "objects"
        ]
        == []
    )
    # empty request is a no-op, not an error
    assert (
        client.post(f"{OS}/buckets/bdx/objects/delete", json={"keys": []}).json()["data"]["deleted"]
        == 0
    )


@mock_aws
def test_batch_delete_and_folder_require_admin(client: TestClient) -> None:
    _wire_moto(client)
    viewer = {"X-CALIBER-User": "@viewer"}
    client.post(f"{OS}/buckets", json={"name": "adx"})
    assert (
        client.post(
            f"{OS}/buckets/adx/objects/delete", json={"keys": ["x"]}, headers=viewer
        ).status_code
        == 403
    )
    assert (
        client.post(f"{OS}/buckets/adx/folders", json={"name": "x"}, headers=viewer).status_code
        == 403
    )
